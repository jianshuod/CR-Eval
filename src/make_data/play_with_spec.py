import re
import os
import docx
import tiktoken
import subprocess
from pathlib import Path
from docx.table import Table
from docx.oxml.table import CT_Tbl
from src.utils.logging import logging
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from src.make_data.play_with_ftp import download_spec
from src.configs import CACHE_CHUNK_DIR, CACHE_IMAGE_DIR
from src.make_data.play_with_docx import getDocxPageCount


logger = logging.getLogger(__name__)


def nodeIter(parent):
    parent_elm = parent.element.body

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def nodeList(parent):
    parent_elm = parent.element.body

    nodeList = []
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            nodeList.append(Paragraph(child, parent))
        elif isinstance(child, CT_Tbl):
            nodeList.append(Table(child, parent))
    return nodeList


def emf2png(src_path, dst_dir, containerID="204c49"):
    os.system(
        " ".join(
            [
                "docker",
                "cp",
                src_path,
                containerID + ":/input/",
                "&&",
                "docker",
                "exec",
                "-it",
                containerID,
                "libreoffice",
                "--headless",
                "--convert-to",
                "png",
                "--outdir",
                "/output",
                "/input/" + src_path.split("/")[-1],
                "&&",
                "docker",
                "cp",
                containerID
                + ":/output/"
                + src_path.split("/")[-1].replace(".emf", ".png"),
                dst_dir,
            ]
        )
    )
    subprocess.run(
        [
            "docker",
            "exec",
            "-it",
            containerID,
            "rm",
            "/input/" + src_path.split("/")[-1],
            "/output/" + src_path.split("/")[-1].replace(".emf", ".png"),
        ]
    )


def getImage(doc, node):
    results = []

    images = node._element.xpath(".//pic:pic")
    if images and len(images) > 0:
        for image in images:
            results += image.xpath(".//a:blip/@r:embed")

    iter = re.finditer('(?<=(v:imagedata r:id=")).*?(?=")', str(node._element.xml))
    results += [i.group() for i in iter]

    if len(results) > 0:
        for r in results:
            rel = doc.part._rels[r]
            imageName = re.findall("/(.*)", rel.target_ref)[0]
            if not os.path.isdir(CACHE_IMAGE_DIR):
                os.mkdir(CACHE_IMAGE_DIR)
            imagePath = os.path.join(CACHE_IMAGE_DIR, imageName)

            imageBlob = rel.target_part.blob
            with open(imagePath, "wb") as imageFile:
                imageFile.write(imageBlob)

            emf2png(imagePath, CACHE_IMAGE_DIR)
            imagePath = imagePath.replace(".emf", ".png")

            # TODO: Query llm
            substitute_text = ""

        return True, substitute_text, imagePath
    else:
        return False, "", ""


def getTable(node: Table):
    table = []

    # FIXME: too slow for tables with many rows
    for id, row in enumerate(node.rows):
        row_list = []
        for cell in row.cells:
            row_list.append(cell.text)
        table.append(row_list)

        # FIXME: Extract only first 50 rows now to accelerate
        if id >= 50:
            break

    if [len(row) for row in table].count(len(table[0])) != len(table):
        return "\n".join([" | ".join(row) for row in table])

    result = ""
    for row in table:
        if result == "":
            result += "|"
            for cell in row:
                result += cell.replace("\n", " ") + "|"
            result += "\n|" + (":-:|" * len(row))
            continue

        result += "\n|"
        for cell in row:
            result += cell.replace("\n", " ") + "|"

    return result


prevStyle = {
    "toc": [0, 0, 0, 0, 0, 0, 0, 0, 0],
    "H": [0, 0, 0, 0, 0, 0, 0, 0, 0],
    "B": [0, 0, 0, 0, 0, 0, 0, 0],
    "Normal": 0,
    "TH": 0,
}


def getParentHeadingID():
    for i in range(7, 0, -1):
        if prevStyle["H"][i] != 0:
            return prevStyle["H"][i]
    if prevStyle["H"][8] != 0:
        return prevStyle["H"][8]
    return prevStyle["H"][0]


def getParentNodeID(id, style: str, text: str):
    if id <= 0:
        return -1

    if style == "TT":
        prevStyle["toc"][7] = id
        prevStyle["toc"][0] = id
        return 0

    elif style.find("toc") == 0 or style.find("H") == 0:
        level = 6
        if style.find(" ") != -1:
            level = int(style.split(" ")[1])
        styleType = "H"
        if style.find("toc") == 0:
            styleType = "toc"
        if level == 1 and prevStyle[styleType][8] != 0:
            return prevStyle[styleType][8]

        prevStyle[styleType][level] = id
        for i in range(level + 1, 7):
            prevStyle[styleType][i] = 0
        return prevStyle[styleType][level - 1]

    elif style.find("B") == 0:
        b_level = int(style[1])
        if text.find("\t") == 0:
            return prevStyle["B"][b_level]
        elif b_level == 1:
            prevStyle["B"][b_level] = id
            return prevStyle["Normal"]

        prevStyle["B"][b_level] = id
        return prevStyle["B"][b_level - 1]

    elif style == "Normal":
        if text.find("\t") == 0:
            return prevStyle["Normal"]

        prevStyle["Normal"] = id
        return getParentHeadingID()

    elif style.find("N") == 0 or style.find("E") == 0:
        return prevStyle["Normal"]

    elif style == "Normal Table":
        if prevStyle["TH"] != 0:
            th = prevStyle["TH"]
            prevStyle["TH"] = 0
            return th
        else:
            return id + 2

    elif style == "TH" and text != "":
        prevStyle["TH"] = id
        return getParentHeadingID()

    elif style == "TH" or style == "TAN" or style == "TF":
        return getParentHeadingID()

    else:
        return -1


# Replace illegal characters for a file name
def clean_file_name(fileName):
    return (
        fileName.replace("\t", " ")
        .replace('"', " ")
        .replace("/", " ")
        .replace("\\", " ")
    )


def build_spec_env(spec_id, spec_version, output_dir):
    # download the spec to the output directory
    local_spec_path = download_spec(spec_id, spec_version, output_dir)

    return SpecContextManager(spec_id, spec_version, local_spec_path)


def num_tokens_from_string(string: str, model_choice="gpt-3.5-turbo") -> int:
    """Returns the number of tokens in a text string."""
    try:
        encoding = tiktoken.encoding_for_model(model_choice)
    except KeyError:
        print("Warning: model not found. Using cl100k_base encoding.")
        encoding = tiktoken.get_encoding("cl100k_base")
    num_tokens = len(encoding.encode(string))
    return num_tokens


def reach_token_limit(text_list, token_limit):
    token = 0
    for s in text_list:
        token += num_tokens_from_string(s)
    return token >= token_limit


class SpecContextManager:
    """
    A context manager for managing a specification at a specific version.
    """

    def __init__(self, spec_id, spec_version, local_spec_path, verbose=False):
        self.spec_id = Path(spec_id).resolve().as_posix()
        self.spec_version = spec_version
        self.local_spec_path = local_spec_path
        self.verbose = verbose

        self.spec_page_count = getDocxPageCount(local_spec_path)
        self.chunked = False
        self.chunk_list = []  # absolute paths to the chunked files

    def process_spec(self, processing_configs):
        pass

    def scrolling_read(self, **kwargs):

        with_heading = kwargs.pop("with_heading", False)
        with_parent = kwargs.pop("with_parent", False)
        discardContext = kwargs.pop("discardContext", False)
        scrollingStride = kwargs.pop("scrollingStride", 3)

        def clickon(text, style):
            return f"**** (ANALYTIC TARGET BEGINS) ****\n{text}\n**** (ANALYTIC TARGET ENDS) ****"

        def scrollrows(table):
            table_rows = [
                " | ".join([cell.text for cell in row.cells]) for row in table
            ]

            if table_rows:
                for i in range(2, len(table_rows), scrollingStride):
                    yield table_rows[i : i + scrollingStride]

        def prepare_parent_trace(style):
            trace = []

            if style.startswith("B") and latestNormal != "":
                trace.append(latestNormal)
                if style != "B1" and latestB1 != "":
                    trace.append(latestB1)
                    if style != "B2" and latestB2 != "":
                        trace.append(latestB2)
                        if style != "B3" and latestB3 != "":
                            trace.append(latestB3)
                            if style != "B4" and latestB4 != "":
                                trace.append(latestB4)

            return trace

        def assemble_final_query(statementList):
            seen = set()
            result = []
            for statement in statementList:
                if statement not in seen:
                    seen.add(statement)
                    result.append(statement)
            return "\n".join(result)

        doc = docx.Document(self.local_spec_path)

        contextUp = ""
        analyticTarget = ""
        contextDown = ""
        latestHeading = ""
        latestNormal = ""
        latestB1 = ""
        latestB2 = ""
        latestB3 = ""
        latestB4 = ""
        latestTableHeading = ""
        analyticTargetStyle = ""
        contextDownStyle = ""
        startFlag = False
        checkInferior = {
            "Normal": "B1",
            "B1": "B2",
            "B2": "B3",
            "B3": "B4",
            "B4": "B5",
            "B5": "B6",
        }

        idx = 0
        currentBuffer = []
        NodeLIST = nodeList(doc)
        while idx < len(NodeLIST):
            node = NodeLIST[idx]
            # Global start
            style = node.style.name
            if not startFlag and not (style.startswith("H") and "General" in node.text):
                idx += 1
                continue
            else:
                startFlag = True

            if style.startswith("H") and "Change history" in node.text:
                break

            if style.startswith(("toc", "TT", "EW", "EX", "TAN")):
                idx += 1
                continue

            if style.startswith("TH"):
                latestTableHeading = node.text
                idx += 1
                continue

            if isinstance(node, Table):
                rowHeader = " | ".join([cell.text for cell in node.rows[0].cells])
                for idx_t, rows in enumerate(scrollrows(node.rows)):
                    contextUp = analyticTarget
                    analyticTarget = contextDown
                    contextDown = "\n".join(rows)
                    if analyticTarget.strip() != "" and idx_t != 0:

                        content_list = []
                        if with_heading:
                            content_list.append(latestHeading)
                            content_list.append(latestTableHeading)
                        content_list.append(rowHeader)
                        if not discardContext:
                            content_list.append(contextUp)
                        content_list.append(
                            clickon(analyticTarget, analyticTargetStyle)
                        )

                        datapoint = assemble_final_query(content_list)

                        yield datapoint
                analyticTargetStyle = style
                contextDownStyle = style
            else:
                # Buffer exploration
                currentBuffer.append(node)
                while len(currentBuffer) < scrollingStride and idx + 1 < len(NodeLIST):
                    nodeX = NodeLIST[idx + 1]
                    styleX = nodeX.style.name
                    if styleX == style:  # Can read in
                        currentBuffer.append(nodeX)
                        idx += 1
                    else:
                        break

                contextUp = analyticTarget
                analyticTarget = contextDown
                contextDown = "\n".join([node_t.text for node_t in currentBuffer])
                analyticTargetStyle = contextDownStyle
                contextDownStyle = style

                inferiorComingFlag = False
                # Check whether the inferior is coming up
                if (
                    "\n" not in analyticTarget
                    and analyticTargetStyle in checkInferior.keys()
                    and checkInferior[analyticTargetStyle] == contextDownStyle
                ):
                    inferiorComingFlag = True

                if (
                    analyticTarget.strip() != ""
                    and not analyticTargetStyle.startswith("H")
                    and not inferiorComingFlag
                ):
                    content_list = []
                    if with_heading:
                        content_list.append(latestHeading)
                    if with_parent:
                        content_list += prepare_parent_trace(analyticTargetStyle)
                    if not discardContext:
                        content_list.append(contextUp)
                    content_list.append(clickon(analyticTarget, analyticTargetStyle))

                    datapoint = assemble_final_query(content_list)
                    yield datapoint

                if style.startswith("H"):
                    latestHeading = currentBuffer[-1].text
                elif style.startswith("B1"):
                    latestB1 = currentBuffer[-1].text
                elif style.startswith("B2"):
                    latestB2 = currentBuffer[-1].text
                elif style.startswith("B3"):
                    latestB3 = currentBuffer[-1].text
                elif style.startswith("B4"):
                    latestB4 = currentBuffer[-1].text
                elif style.startswith("Normal"):
                    latestNormal = currentBuffer[-1].text

            currentBuffer = []
            idx += 1

    def apply_chunks(self, chunking_mode, **kwargs):
        if self.chunked:
            return

        if chunking_mode == "section":
            doc = docx.Document(self.local_spec_path)

            node_hierarchy = {}
            nodeid2title = {}
            nodeid2path = {}
            chunk_list = {}
            cur_heading = 0

            for id, node in enumerate(nodeIter(doc)):
                nodeID = id
                style = node.style.name
                if isinstance(node, Table):
                    text = getTable(node)
                else:
                    text = node.text.replace("\n", "\\n")
                parentNodeID = getParentNodeID(nodeID, style, text)

                if style == "TH":
                    isFind, sub_text, path = getImage(doc, node)
                    if isFind:
                        text = f"![{sub_text}]({path})"

                node_hierarchy[nodeID] = parentNodeID

                if style.find("H") == 0:
                    cur_heading = nodeID
                    nodeid2title[nodeID] = text
                    nodeid2path[nodeID] = Path(CACHE_CHUNK_DIR)
                else:
                    if cur_heading not in chunk_list:
                        chunk_list[cur_heading] = [text]
                    else:
                        chunk_list[cur_heading].append(text)

                def get_position_trace(nodeID):
                    position_trace = ""
                    while nodeID in node_hierarchy:
                        if nodeID in nodeid2title:
                            position_trace = (
                                nodeid2title[nodeID] + " > " + position_trace
                            )
                        nodeID = node_hierarchy[nodeID]  # to the parent
                    return position_trace

            for nodeID in chunk_list:
                abs_path = os.path.join(
                    CACHE_CHUNK_DIR,
                    clean_file_name(nodeid2title.get(nodeID, str(nodeID))),
                )

                # Save chunk content to local storage
                chunk_file = open(
                    abs_path,
                    "w",
                    encoding="utf-8",
                )
                chunk_file.write(get_position_trace(nodeID) + "\n")
                for sentence in chunk_list[nodeID]:
                    chunk_file.write(sentence + "\n")
                chunk_file.close()

                self.chunk_list.append(abs_path)

            self.chunked = True

        elif chunking_mode == "block":
            doc = docx.Document(self.local_spec_path)

            node_hierarchy = {}
            nodeid2title = {}
            nodeid2path = {}
            chunk_list = {}
            cur_heading = 0

            for id, node in enumerate(nodeIter(doc)):
                nodeID = id
                style = node.style.name
                if isinstance(node, Table):
                    text = getTable(node)
                else:
                    text = node.text.replace("\n", "\\n")
                parentNodeID = getParentNodeID(nodeID, style, text)

                if style == "TH":
                    isFind, sub_text, path = getImage(doc, node)
                    if isFind:
                        text = f"![{sub_text}]({path})"

                node_hierarchy[nodeID] = parentNodeID

                if style.find("H") == 0:
                    cur_heading = nodeID
                    nodeid2title[nodeID] = text
                    nodeid2path[nodeID] = Path(CACHE_CHUNK_DIR)
                else:
                    if cur_heading not in chunk_list:
                        chunk_list[cur_heading] = [[text]]
                    elif style.find("Normal") == 0:
                        if not reach_token_limit(chunk_list[cur_heading][-1], 1000):
                            chunk_list[cur_heading][-1].append(text)
                        else:
                            chunk_list[cur_heading].append([text])
                    else:
                        chunk_list[cur_heading][-1].append(text)

                def get_position_trace(nodeID):
                    position_trace = ""
                    while nodeID in node_hierarchy:
                        if nodeID in nodeid2title:
                            position_trace = (
                                nodeid2title[nodeID] + " > " + position_trace
                            )
                        nodeID = node_hierarchy[nodeID]  # to the parent
                    return position_trace

            for nodeID in chunk_list:
                for id, block in enumerate(chunk_list[nodeID]):
                    clean_name = str(
                        clean_file_name(nodeid2title.get(nodeID, str(nodeID)))
                    )
                    file_name = (
                        clean_name.replace(" ", "-" + str(id) + " ", 1)
                        if clean_name.find(" ") != -1
                        else clean_name + "-" + str(id)
                    )
                    abs_path = os.path.join(CACHE_CHUNK_DIR, file_name)

                    # Save chunk content to local storage
                    chunk_file = open(abs_path, "w", encoding="utf-8")
                    chunk_file.write(get_position_trace(nodeID) + str(id) + "\n" * 2)
                    for sentence in block:
                        chunk_file.write(sentence + "\n")
                    chunk_file.close()

                    self.chunk_list.append(abs_path)

            self.chunked = True

        elif chunking_mode == "size-aware":
            # load a tokenizer
            from ..analysis.utils import num_tokens_from_string

            chunk_limit = kwargs.get("chunk_limit", 2000)

            doc = docx.Document(self.local_spec_path)

            chunk_list = []

            chunkBuffer = ""
            bufferSize = 0

            startFlag = False

            for id, node in enumerate(nodeIter(doc)):
                style = node.style.name

                if not startFlag and not (
                    style.startswith("H") and "General" in node.text
                ):
                    continue
                else:
                    startFlag = True

                if isinstance(node, Table):
                    continue

                text = node.text.replace("\n", "\\n")

                textSize = num_tokens_from_string("gpt-3.5-turbo", text)

                if bufferSize + textSize > chunk_limit:
                    chunk_list.append(chunkBuffer)
                    chunkBuffer = ""
                    bufferSize = 0

                chunkBuffer += text + "\n\n"
                bufferSize += textSize

            if chunkBuffer:
                chunk_list.append(chunkBuffer)

            self.chunk_list = chunk_list

            self.chunked = True

        else:
            raise NotImplementedError(
                f"Chunking mode {chunking_mode} is not implemented yet."
            )

    def clear_chunks(self):
        if not self.chunked:
            return

        # clear chunks
        for chunk_path in self.chunk_list:
            os.remove(chunk_path)
        self.chunked = False
        self.chunk_list.clear()

    def list_chunks(self):
        logger.info(f"Loading chunk list... Totally, {len(self.chunk_list)} chunks.")
        return self.chunk_list
