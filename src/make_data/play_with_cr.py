import json
import docx
from typing import Dict
from docx.table import Table
from pathlib import Path, PosixPath
from src.utils.logging import logging
from docx.text.paragraph import Paragraph
from src.make_data.instance import ChangeRequest
from src.make_data.play_with_spec import nodeList
from src.make_data.play_with_docx import (
    get_cell_text,
    getDocxPageCount,
    count_images,
    detect_table_modification,
    deny_all_text,
    accept_all_text,
    accept_all_table,
    deny_all_table,
    print_xml_structure,
)

logger = logging.getLogger(__name__)


"""Extracts the fields from the CR document.

Args:
    cr_path: the local path to the CR document
    output_dir: the directory to save the extracted fields

Return:
    a dict of the CR key elements, including
        - spec: the target specification number √
        - current version: √
        - meeting id: √
        - title: √
        - reason_for_change: √
        - summary_of_change: √
        - consequences_if_not_approved: √
        - clauses_affected: √
        - other_specs_affected: √
        - change_list: √
        - extracted_index: √
        - input_index: √
        - date: √
        - table_modified_flag: √
        - page_count: √
        - figure_modified_flag: √
        - other_comments: √

Note that this function assumes a standard 3GPP CR document format.
"""


def get_cr_fields(cr_path: PosixPath, output_dir, include_table=False):

    if cr_path.suffix == ".json":
        with open(cr_path, "r", encoding="utf-8") as json_file:
            cr = ChangeRequest(**json.load(json_file))
            return cr
    elif cr_path.suffix == ".docx":
        return ChangeRequest(
            **get_cr_fields_relative_mode(cr_path, output_dir, include_table)
        )
    else:
        raise ValueError("Unsupported file format.")


def is_cr(doc: docx.Document):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "CHANGE REQUEST" in get_cell_text(cell):
                    return True
    return False


def get_cr_fields_relative_mode(cr_path, output_dir, include_table=False) -> Dict:

    cr_path = Path(cr_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    file_prefix = cr_path.stem

    doc = docx.Document(cr_path)
    if not doc or not is_cr(doc):
        return None

    other_specs_affected_o_and_m = None
    other_specs_affected_test = None
    other_specs_affected_core = None
    source_to_tsg = None
    source_to_wg = None
    source = None

    last_visited = None
    for tid, table in enumerate(doc.tables):
        for rid, row in enumerate(table.rows):
            for cid, cell in enumerate(row.cells):
                cell_text: str = get_cell_text(cell).strip()
                if last_visited is not None:
                    if "CR" == cell_text:
                        spec = last_visited
                    elif "SpecNumber" in last_visited:
                        spec = cell_text
                    elif "Current version:" in last_visited:
                        current_version = cell_text
                    elif "Title:" in last_visited:
                        title = cell_text
                    elif "Source" in last_visited:  # legacy version
                        source = cell_text
                    elif "Source to WG" in last_visited:
                        source_to_wg = cell_text
                    elif "Source to TSG" in last_visited:
                        source_to_tsg = cell_text
                    elif "Work item code" in last_visited:
                        work_item_code = cell_text
                    elif "Category:" in last_visited:
                        category = cell_text
                    elif "Date:" in last_visited:
                        date = cell_text
                    elif "Reason for change:" in last_visited:
                        reason_for_change = cell_text
                    elif "Summary of change:" in last_visited:
                        summary_of_change = cell_text
                    elif "Consequences" in last_visited:
                        consequences_if_not_approved = cell_text
                    elif "Clauses affected:" in last_visited:
                        clauses_affected = cell_text.split("; ")
                    elif "other core specifications" in last_visited.lower():
                        other_specs_affected_core = cell_text
                    elif "Test specifications" in last_visited:
                        other_specs_affected_test = cell_text
                    elif "O&M Specifications" in last_visited:
                        other_specs_affected_o_and_m = cell_text
                    elif "Other comments" in last_visited:
                        other_comments = cell_text
                last_visited = cell_text

    other_specs_affected = ""
    if (
        other_specs_affected_core is not None
        and "TS/TR ... CR ..." not in other_specs_affected_core
        and other_specs_affected_core != ""
    ):
        other_specs_affected += other_specs_affected_core
    if (
        other_specs_affected_test is not None
        and "TS/TR ... CR ..." not in other_specs_affected_test
        and other_specs_affected_test != ""
    ):
        other_specs_affected += other_specs_affected_test
    if (
        other_specs_affected_o_and_m is not None
        and "TS/TR ... CR ..." not in other_specs_affected_o_and_m
        and other_specs_affected_o_and_m != ""
    ):
        other_specs_affected += other_specs_affected_o_and_m

    other_specs_affected = other_specs_affected.split("; ")

    # extract meeting id and extracted_index
    i = 0
    while not doc.paragraphs[i].text.strip():
        i += 1
    meeting_id, extracted_index = doc.paragraphs[i].text.split("\t", 1)

    # extract change list
    changes = ()
    raw_version = ""
    changed_version = ""
    for nid, node in enumerate(nodeList(doc)[7:]):
        try:
            if isinstance(node, Paragraph):
                raw_version += f"{deny_all_text(node)}\n"
                changed_version += f"{accept_all_text(node)}\n"
            elif isinstance(node, Table) and include_table:
                raw_version += f"{deny_all_table(node)}\n"
                changed_version += f"{accept_all_table(node)}\n"
        except Exception as e:
            logger.warning(f"Failed to process node {nid}: {str(e)}")
            print_xml_structure(node._element)
            continue
    if raw_version or changed_version:
        changes = (raw_version, changed_version)

    # artifact modified
    table_modified_flag = detect_table_modification(doc)
    page_count = getDocxPageCount(cr_path)
    figure_count = count_images(doc)
    figure_modified_flag = figure_count > 0

    # establish the fields
    fields = {
        "spec": spec,
        "current_version": current_version,
        "meeting_id": meeting_id,
        "title": title,
        "reason_for_change": reason_for_change,
        "summary_of_change": summary_of_change,
        "consequences_if_not_approved": consequences_if_not_approved,
        "clauses_affected": clauses_affected,
        "other_specs_affected": other_specs_affected,
        "extracted_index": extracted_index,
        "date": date,
        "change_list": changes,
        "table_modified_flag": table_modified_flag,
        "page_count": page_count,
        "figure_modified_flag": figure_modified_flag,
        "other_comments": other_comments,
        "source": source,
        "source_to_wg": source_to_wg,
        "source_to_tsg": source_to_tsg,
        "work_item_code": work_item_code,
        "category": category,
    }

    # Save to local directory
    attr = (
        f"t{'t' if table_modified_flag else 'n'}{'f' if figure_modified_flag else 'n'}"
    )
    output_path = output_dir / f"{attr}-{file_prefix}.json"

    with open(output_path, "w", encoding="utf-8") as json_file:
        json.dump(fields, json_file, ensure_ascii=False, indent=4)

    # logger.info(f"Fields of {extracted_index} saved to {output_path}")

    return fields
