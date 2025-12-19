import os
import re
import docx
import json
import logging
import zipfile
import subprocess
from pathlib import Path
from docx.oxml.ns import qn, nsmap


logging.basicConfig(
    filename="extract_cr_docs.log",
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.get  # logger(__name__)

failure_cases = []


def getDocxPageCount(docx_fpath):
    # logger.info(f"Getting page count of {docx_fpath}")
    docx_object = zipfile.ZipFile(docx_fpath)
    docx_property_file_data = docx_object.read("docProps/app.xml").decode()
    page_count = re.search(r"<Pages>(\d+)</Pages>", docx_property_file_data).group(1)
    return int(page_count)


def get_cell_text(cell):
    def _get_cell_text(element):
        text = ""
        for child in element:
            if child.tag == qn("w:t"):
                text += child.text or ""
            else:
                text += _get_cell_text(child)
        return text

    nsmap["mc"] = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    return _get_cell_text(cell._element)


def print_xml_structure(element, indent=0):
    spaces = " " * indent
    print(f"{spaces}{element.tag}: {element.text}")

    for child in element:
        print_xml_structure(child, indent + 2)


def accepted_text(p):
    def _accepted_text(p):
        text = ""
        for run in p.xpath("w:r[not(w:pPr/w:rPr/w:moveFrom)] | w:ins/w:r"):
            for child in run:
                if child.tag == qn("w:t"):
                    text += child.text or ""
                elif child.tag == qn("w:tab"):
                    text += "\t"
                elif child.tag in (qn("w:br"), qn("w:cr")):
                    text += "\n"
                elif child.tag == qn("mc:AlternateContent"):
                    for nested_p in child.xpath("mc:Choice[1]//w:p", namespaces=nsmap):
                        text += _accepted_text(nested_p)
                        text += "\n"
        return text

    nsmap["mc"] = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    return _accepted_text(p._p)


def display_row(row):
    for idx, cell in enumerate(row.cells):
        print(f"cell {idx}: {cell.text}")


def get_cr_fields(cr_path, output_dir):
    """
    Extracts the fields from the CR document.

    Args:
        cr_path: the local path to the CR document
        output_dir: the directory to save the extracted fields

    Return:
        an dict of the CR key elements, including
            - spec: the target specification number √
            - current version: √
            - meeting id: √
            - title: √
            - reason_for_change: √
            - summary_of_change: √
            - consequences_if_not_approved: √
            - clauses_affected: √
            - other_specs_affected: √
            - change_list: √
            - extracted_index: √
            - input_index: √
            - date: √
            - table_modified_flag: √
            - page_count: √
            - figure_modified_flag: √

    Note that this function assumes a standard 3GPP CR document format.
    """
    os.makedirs(output_dir, exist_ok=True)

    doc = docx.Document(cr_path)
    if not doc:
        # logger.error(f"Failed to open {cr_path}")
        return fields

    # logger.info(f"Extracting fields from {cr_path}")
    # extract meeting id and extracted_index
    meeting_id, extracted_index = doc.paragraphs[0].text.split("\t", 1)

    #
    spec = get_cell_text(doc.tables[0].cell(3, 1))
    current_version = get_cell_text(doc.tables[0].cell(3, 7))
    date = get_cell_text(doc.tables[2].cell(6, 10))

    # extract information
    title = doc.tables[2].cell(1, 1).text
    reason_for_change = doc.tables[2].cell(11, 2).text
    summary_of_change = doc.tables[2].cell(13, 2).text
    consequences_if_not_approved = doc.tables[2].cell(15, 2).text
    clauses_affected = doc.tables[2].cell(17, 2).text.split("; ")
    other_specs_affected_core = doc.tables[2].cell(20, 8).text
    other_specs_affected_test = doc.tables[2].cell(21, 8).text
    other_specs_affected_o_and_m = doc.tables[2].cell(22, 8).text

    other_specs_affected = []
    if "TS/TR ... CR ..." not in other_specs_affected_core:
        other_specs_affected += [other_specs_affected_core]
    if "TS/TR ... CR ..." not in other_specs_affected_test:
        other_specs_affected += [other_specs_affected_test]
    if "TS/TR ... CR ..." not in other_specs_affected_o_and_m:
        other_specs_affected += [other_specs_affected_o_and_m]

    # extract change list
    changes = []
    raw_version = ""
    changed_version = ""
    for p in doc.paragraphs:
        if "change" in p.text.lower():  # control line
            if "start" in p.text.lower() or "first" in p.text.lower():
                pass
            elif "end" in p.text.lower() or "next" in p.text.lower():
                changes.append((raw_version, changed_version))
            raw_version = ""
            changed_version = ""
        else:  # continue
            raw_version += f"{p.text}\n"
            changed_version += f"{accepted_text(p)}\n"

    # table modified
    table_modified_flag = len(doc.tables) > 3
    page_count = getDocxPageCount(cr_path)
    figure_count = len([p for p in doc.paragraphs if "Figure" in p.text])
    figure_modified_flag = figure_count > 0

    # establish the fields
    fields = {
        "spec": spec,
        "current_version": current_version,
        "meeting_id": meeting_id,
        "title": title,
        "reason_for_change": reason_for_change,
        "summary_of_change": summary_of_change,
        "consequences_if_not_approved": consequences_if_not_approved,
        "clauses_affected": clauses_affected,
        "other_specs_affected": other_specs_affected,
        "extracted_index": extracted_index,
        "date": date,
        "change_list": changes,
        "table_modified_flag": table_modified_flag,
        "page_count": page_count,
        "figure_modified_flag": figure_modified_flag,
    }

    # Save to local directory
    output_path = os.path.join(output_dir, f"{fields['input_index']}.json")
    with open(output_path, "w", encoding="utf-8") as json_file:
        json.dump(fields, json_file, ensure_ascii=False, indent=4)

    # logger.info(f"Fields of {fields['extracted_index']} saved to {output_path}")

    return fields


def convert_doc_to_docx(doc_path):
    # 确保文件路径是Path对象
    doc_path = Path(doc_path)

    # 检查文件是否已经是.docx格式
    if doc_path.suffix == ".docx":
        return doc_path

    # 构建输出文件的路径（原地转换意味着同一目录）
    docx_path = doc_path.with_suffix(".docx")

    # 调用LibreOffice命令行进行转换
    try:
        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(doc_path.parent),
                str(doc_path),
            ],
            check=True,
        )
        # logger.info(f'Converted: {doc_path} to {docx_path}')

        # 如果转换成功，删除原.doc文件
        doc_path.unlink()
    except subprocess.CalledProcessError as e:
        # logger.warning(f'Error during conversion: {e}')
        #
        pass
    return docx_path


def unzip_file(zip_path, extract_to):
    extract_to = Path(extract_to)
    extract_to.mkdir(parents=True, exist_ok=True)

    if not zipfile.is_zipfile(zip_path):
        logging.error(f"{zip_path} is not a valid zip file.")
        return None

    with zipfile.ZipFile(zip_path, "r") as zip_ref:
        doc_files = [
            f for f in zip_ref.namelist() if f.endswith(".doc") or f.endswith(".docx")
        ]
        for file in doc_files:
            try:
                zip_ref.extract(file, extract_to)
                original_path = extract_to / file
                new_path = extract_to / Path(file).name
                original_path.rename(new_path)
                logging.info(f"Extracted {file} to {new_path}")
            except zipfile.BadZipFile as e:
                logging.error(f"Corrupt file {file} in {zip_path}: {e}")
                continue

        if len(doc_files) == 1:
            return extract_to / Path(doc_files[0]).name
        elif len(doc_files) > 1:
            logging.info("More than one .doc or .docx file extracted")
            return None
        else:
            logging.info("No .doc or .docx files found in the archive")
            return None


# Enhance error handling in process_cr_zip
def process_cr_zip(cr_path, target_dir):
    try:
        target_dir = Path(target_dir)
        target_doc_dir = target_dir / "docs"
        target_json_dir = target_dir / "jsons"
        docx_path = unzip_file(cr_path, target_doc_dir)
        if docx_path is None or not docx_path.suffix == ".docx":
            logging.error(f"No valid .docx file found in {cr_path}")
            failure_cases.append(cr_path)
            return

        get_cr_fields(docx_path, target_json_dir)
    except Exception as e:
        logging.error(f"Failed to process {cr_path}: {str(e)}")
        failure_cases.append(cr_path)


def save_failure_cases(file_path):
    """Save the failure cases to a text file."""
    with open(file_path, "w") as file:
        for case in failure_cases:
            file.write(case + "\n")
